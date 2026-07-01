# ============================================================
# text_extraction.py — PDF & DOCX Text Extraction
# ============================================================
# Handles parsing uploaded resume files into plain text.
# Supports PDF (via PyMuPDF/fitz) and DOCX (via python-docx).
# ============================================================

import io
import fitz  # PyMuPDF
from docx import Document


def extract_text_from_pdf(uploaded_file: object) -> str:
    """
    Extract all text content from a PDF file.

    Uses PyMuPDF (fitz) for fast, reliable text extraction.
    Handles multi-page PDFs by concatenating all page text.

    Args:
        uploaded_file: A Streamlit UploadedFile object (PDF).

    Returns:
        Extracted plain text string, or empty string on failure.
    """
    try:
        # Read the uploaded file bytes into a PyMuPDF document
        pdf_bytes = uploaded_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        text_parts = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_text = page.get_text("text")
            if page_text:
                text_parts.append(page_text)

        doc.close()
        return "\n".join(text_parts).strip()

    except Exception as e:
        print(f"[ERROR] Failed to extract text from PDF '{uploaded_file.name}': {e}")
        return ""


def extract_text_from_docx(uploaded_file: object) -> str:
    """
    Extract all text content from a DOCX file.

    Uses python-docx to iterate over all paragraphs and tables,
    capturing text that might be stored in table cells (common
    in formatted resumes).

    Args:
        uploaded_file: A Streamlit UploadedFile object (DOCX).

    Returns:
        Extracted plain text string, or empty string on failure.
    """
    try:
        # Read bytes into a BytesIO stream for python-docx
        docx_bytes = uploaded_file.read()
        doc = Document(io.BytesIO(docx_bytes))

        text_parts = []

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Also extract text from tables (many resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in text_parts:
                        text_parts.append(cell_text)

        return "\n".join(text_parts).strip()

    except Exception as e:
        print(f"[ERROR] Failed to extract text from DOCX '{uploaded_file.name}': {e}")
        return ""


def extract_text(uploaded_file: object) -> str:
    """
    Router function — detects file type and calls the appropriate extractor.

    Args:
        uploaded_file: A Streamlit UploadedFile object.

    Returns:
        Extracted text string. Returns empty string if:
        - File type is unsupported
        - Extraction fails
        - File is empty/corrupted
    """
    if uploaded_file is None:
        return ""

    file_name = uploaded_file.name.lower()

    # Reset file pointer to beginning (in case it was read before)
    uploaded_file.seek(0)

    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif file_name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    else:
        print(f"[WARNING] Unsupported file type: {file_name}")
        return ""
