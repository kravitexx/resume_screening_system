import io
import datetime
import markdown
from docx import Document
from htmldocx import HtmlToDocx
from fpdf import FPDF

def generate_markdown_report(results: list[dict], is_pro: bool, compare_summary: str = None) -> str:
    """Generates a Markdown string from the analysis results."""
    lines = []
    lines.append("# ResumeAI Analysis Report")
    lines.append(f"**Generated on:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"**Mode:** {'PRO (Semantic AI)' if is_pro else 'BASIC (Keyword)'}")
    lines.append("---")
    
    if compare_summary:
        lines.append("## AI Comparison Summary")
        lines.append(compare_summary)
        lines.append("---")
        
    for r in results:
        lines.append(f"## Candidate: {r['file_name']}")
        if is_pro:
            lines.append(f"- **Semantic Score:** {r['semantic_score']}%")
            lines.append(f"- **Fit Level:** {r['fit_level']}")
            lines.append(f"- **Fit Assessment:** {r['fit_assessment']}")
            lines.append(f"- **Matched Skills:** {', '.join(r['matched_skills']) if r.get('matched_skills') else 'None'}")
            lines.append(f"- **Missing Skills:** {', '.join(r['missing_skills']) if r.get('missing_skills') else 'None'}")
        else:
            lines.append(f"- **Composite Score:** {r['composite_score']}%")
            lines.append(f"- **ATS Score:** {r['ats_score']}%")
            lines.append(f"- **Keyword Match:** {r['keyword_match_count']}/{r['keyword_total']}")
            lines.append(f"- **Sections Score:** {r['section_score']}%")
            lines.append(f"- **Matched Keywords:** {', '.join(r['matched_keywords']) if r.get('matched_keywords') else 'None'}")
            lines.append(f"- **Missing Keywords:** {', '.join(r['missing_keywords']) if r.get('missing_keywords') else 'None'}")
            contact = r.get('contact_info', {})
            lines.append(f"- **Contact:** Email: {contact.get('email', 'N/A')} | Phone: {contact.get('phone', 'N/A')}")
        lines.append("\n---\n")
        
    return "\n".join(lines)


def generate_docx_report(results: list[dict], is_pro: bool, compare_summary: str = None) -> io.BytesIO:
    """Generates a DOCX file from the analysis results."""
    doc = Document()
    doc.add_heading('ResumeAI Analysis Report', 0)
    
    p = doc.add_paragraph()
    p.add_run(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = True
    p.add_run(f"Mode: {'PRO (Semantic AI)' if is_pro else 'BASIC (Keyword)'}").bold = True
    
    md_body = []
    if compare_summary:
        md_body.append("## AI Comparison Summary")
        md_body.append(compare_summary)
        md_body.append("---")
        
    for r in results:
        md_body.append(f"## Candidate: {r['file_name']}")
        if is_pro:
            md_body.append(f"- **Semantic Score:** {r['semantic_score']}%")
            md_body.append(f"- **Fit Level:** {r['fit_level']}")
            md_body.append(f"- **Fit Assessment:** {r['fit_assessment']}")
            md_body.append(f"- **Matched Skills:** {', '.join(r['matched_skills']) if r.get('matched_skills') else 'None'}")
            md_body.append(f"- **Missing Skills:** {', '.join(r['missing_skills']) if r.get('missing_skills') else 'None'}")
        else:
            md_body.append(f"- **Composite Score:** {r['composite_score']}%")
            md_body.append(f"- **ATS Score:** {r['ats_score']}%")
            md_body.append(f"- **Keyword Match:** {r['keyword_match_count']}/{r['keyword_total']}")
            md_body.append(f"- **Sections Score:** {r['section_score']}%")
            md_body.append(f"- **Matched Keywords:** {', '.join(r['matched_keywords']) if r.get('matched_keywords') else 'None'}")
            md_body.append(f"- **Missing Keywords:** {', '.join(r['missing_keywords']) if r.get('missing_keywords') else 'None'}")
            contact = r.get('contact_info', {})
            md_body.append(f"- **Contact:** Email: {contact.get('email', 'N/A')} | Phone: {contact.get('phone', 'N/A')}")
        md_body.append("\n---\n")
        
    final_md = "\n".join(md_body)
    
    # Parse Markdown into HTML
    html_text = markdown.markdown(final_md)
    
    # Inject HTML directly into DOCX
    new_parser = HtmlToDocx()
    new_parser.add_html_to_document(html_text, doc)
    
    # Save to BytesIO
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f


class PDFReport(FPDF):
    def header(self):
        self.set_font("helvetica", "B", 15)
        self.set_text_color(20, 30, 50)
        self.cell(0, 10, "ResumeAI Analysis Report", border=False, ln=True, align="C")
        self.set_font("helvetica", "I", 10)
        self.set_text_color(100, 100, 100)
        self.cell(0, 8, f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", border=False, ln=True, align="C")
        self.set_draw_color(200, 200, 200)
        self.line(10, 28, 200, 28)
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font("helvetica", "I", 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def generate_pdf_report(results: list[dict], is_pro: bool, compare_summary: str = None) -> bytes:
    """Generates a PDF file from the analysis results using HTML parsing."""
    def sanitize(text):
        if not text: return ""
        # Remove characters that can't map to latin-1
        return text.encode('latin-1', 'replace').decode('latin-1')
        
    pdf = PDFReport()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    pdf.set_font("helvetica", "B", 12)
    pdf.set_text_color(50, 100, 200)
    pdf.cell(0, 10, f"Mode: {'PRO (Semantic AI)' if is_pro else 'BASIC (Keyword)'}", ln=True)
    pdf.ln(5)
    
    md_body = []
    if compare_summary:
        md_body.append("<h2>AI Comparison Summary</h2>")
        # Convert the compare_summary markdown to HTML
        md_body.append(markdown.markdown(compare_summary))
        md_body.append("<hr>")
        
    for r in results:
        md_body.append(f"<h2>Candidate: {r['file_name']}</h2>")
        if is_pro:
            md_body.append(
                f"<ul>"
                f"<li><b>Semantic Score:</b> {r['semantic_score']}%</li>"
                f"<li><b>Fit Level:</b> {r['fit_level']}</li>"
                f"<li><b>Fit Assessment:</b> {r['fit_assessment']}</li>"
                f"<li><b>Matched Skills:</b> {', '.join(r['matched_skills']) if r.get('matched_skills') else 'None'}</li>"
                f"<li><b>Missing Skills:</b> {', '.join(r['missing_skills']) if r.get('missing_skills') else 'None'}</li>"
                f"</ul>"
            )
        else:
            contact = r.get('contact_info', {})
            md_body.append(
                f"<ul>"
                f"<li><b>Composite Score:</b> {r['composite_score']}%</li>"
                f"<li><b>ATS Score:</b> {r['ats_score']}%</li>"
                f"<li><b>Keyword Match:</b> {r['keyword_match_count']}/{r['keyword_total']}</li>"
                f"<li><b>Sections Score:</b> {r['section_score']}%</li>"
                f"<li><b>Matched Keywords:</b> {', '.join(r['matched_keywords']) if r.get('matched_keywords') else 'None'}</li>"
                f"<li><b>Missing Keywords:</b> {', '.join(r['missing_keywords']) if r.get('missing_keywords') else 'None'}</li>"
                f"<li><b>Contact:</b> Email: {contact.get('email', 'N/A')} | Phone: {contact.get('phone', 'N/A')}</li>"
                f"</ul>"
            )
        md_body.append("<hr>")
        
    final_html = "".join(md_body)
    sanitized_html = sanitize(final_html)
    
    pdf.set_text_color(0, 0, 0)
    pdf.write_html(sanitized_html)
    
    return bytes(pdf.output())
